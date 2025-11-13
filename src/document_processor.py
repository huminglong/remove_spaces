"""
文件名: src/document_processor.py
功能描述: Word文档处理器，负责读取、分析和修改Word文档内容，保持文档格式不变
主要函数:
  - load_document(): 加载Word文档
  - get_all_text(): 获取文档中的所有文本内容
  - update_text_content(): 更新文档中的文本内容
  - save_document(): 保存文档到指定路径
  - save_as_new_document(): 创建新文档并保存处理后的文本
  - get_document_info(): 获取文档基本信息
主要类:
  - DocumentProcessor: Word文档处理核心类
"""

import docx
from typing import List, Dict, Tuple, Optional
import re
from pathlib import Path

from src.exceptions import (
    DocumentLoadError,
    DocumentStructureError,
    TextUpdateError,
    DocumentSaveError
)
from src.logger_config import setup_logger
from config.settings import settings


class DocumentProcessor:
    """
    Word文档处理器，负责读取和写入Word文档

    该类提供Word文档的加载、文本提取、内容更新和保存功能。
    重要原则：只修改文字内容，严格保持文档的所有格式、样式、图片等非文字元素不变。

    Attributes:
        document: 加载的Word文档对象
        original_structure: 文档的原始结构信息
    """

    def __init__(self):
        """
        初始化文档处理器

        创建DocumentProcessor实例，初始化文档和结构存储。
        """
        self.document = None
        self.original_structure = []
        self.logger = setup_logger(__name__)
    
    def __enter__(self):
        """
        上下文管理器入口
        
        支持with语句，方便资源管理。
        
        Returns:
            DocumentProcessor: 返回self以支持with语句
            
        Examples:
            >>> with DocumentProcessor() as processor:
            ...     processor.load_document("test.docx")
            ...     # 自动清理资源
        """
        return self
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        """
        上下文管理器出口
        
        确保资源被正确释放。
        
        Args:
            exc_type: 异常类型
            exc_val: 异常值
            exc_tb: 异常追踪信息
            
        Returns:
            bool: False表示不抑制异常
        """
        self.close()
        return False
    
    def close(self):
        """
        关闭文档并释放资源
        
        清理文档对象和结构信息，释放内存。
        
        Examples:
            >>> processor = DocumentProcessor()
            >>> processor.load_document("test.docx")
            >>> processor.close()
        """
        if self.document:
            self.logger.debug("关闭文档处理器，释放资源")
            self.document = None
            self.original_structure = []
    
    def load_document(self, file_path: str) -> bool:
        """
        加载Word文档
        
        执行完整的输入验证,包括文件存在性、格式、大小等检查。
        
        Args:
            file_path: 文档路径
            
        Returns:
            bool: 加载成功返回True，否则返回False
            
        Raises:
            DocumentLoadError: 当文档加载失败时抛出
            
        Examples:
            >>> processor = DocumentProcessor()
            >>> processor.load_document("test.docx")
            True
        """
        # 输入验证
        if not file_path:
            self.logger.error("文件路径不能为空")
            raise ValueError("文件路径不能为空")
        
        if not isinstance(file_path, (str, Path)):
            self.logger.error(f"文件路径必须是字符串或Path对象，实际类型: {type(file_path)}")
            raise TypeError(f"文件路径必须是字符串或Path对象，实际类型: {type(file_path)}")
        
        file_path = Path(file_path)
        
        # 文件存在性检查
        if not file_path.exists():
            self.logger.error(f"文件不存在: {file_path}")
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 文件格式检查
        if file_path.suffix.lower() not in settings.DocumentProcessing.SUPPORTED_EXTENSIONS:
            self.logger.error(f"不支持的文件格式: {file_path.suffix}，仅支持 {settings.DocumentProcessing.SUPPORTED_EXTENSIONS}")
            raise ValueError(f"不支持的文件格式: {file_path.suffix}")
        
        # 文件大小检查
        file_size = file_path.stat().st_size
        if file_size > settings.DocumentProcessing.MAX_FILE_SIZE:
            self.logger.error(f"文件过大: {file_size} 字节，最大支持 {settings.DocumentProcessing.MAX_FILE_SIZE} 字节")
            raise ValueError(f"文件过大，最大支持 {settings.DocumentProcessing.MAX_FILE_SIZE / (1024*1024):.1f}MB")
        
        try:
            self.logger.info(f"开始加载文档: {file_path}")
            self.document = docx.Document(str(file_path))
            self._extract_document_structure()
            self.logger.info(f"成功加载文档: {file_path}，包含 {len(self.original_structure)} 个元素")
            return True
        except FileNotFoundError:
            raise
        except ValueError:
            raise
        except PermissionError as e:
            self.logger.error(f"权限不足，无法访问文件: {file_path}", exc_info=True)
            raise PermissionError(f"权限不足，无法访问文件: {file_path}") from e
        except Exception as e:
            self.logger.error(f"加载文档失败: {file_path}", exc_info=True)
            raise DocumentLoadError(f"加载文档失败: {str(e)}") from e
    
    def _extract_document_structure(self) -> None:
        """
        提取文档的完整结构，包括所有格式信息

        遍历文档的所有段落和表格，保存其原始结构、文本内容和格式信息，
        为后续的文本更新操作提供基础。
        """
        self.original_structure = []

        # 提取段落及其完整结构
        for para in self.document.paragraphs:
            if para.text.strip():  # 只处理非空段落
                self.original_structure.append({
                    'type': 'paragraph',
                    'paragraph': para,  # 保存段落对象引用
                    'original_text': para.text,
                    'runs': self._extract_runs(para)
                })

        # 提取表格及其完整结构
        for table in self.document.tables:
            table_structure = {
                'type': 'table',
                'table': table,  # 保存表格对象引用
                'rows': []
            }

            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_data = {
                        'cell': cell,  # 保存单元格对象引用
                        'original_text': cell.text,
                        'paragraphs': []
                    }

                    # 提取单元格中的所有段落
                    for para in cell.paragraphs:
                        if para.text.strip():
                            cell_data['paragraphs'].append({
                                'paragraph': para,  # 保存段落对象引用
                                'original_text': para.text,
                                'runs': self._extract_runs(para)
                            })

                    row_data.append(cell_data)

                if row_data:
                    table_structure['rows'].append(row_data)

            if table_structure['rows']:
                self.original_structure.append(table_structure)
    
    def _extract_runs(self, paragraph) -> List[Dict]:
        """
        提取段落中的runs（文本片段）及其完整格式信息

        遍历段落中的所有run对象，提取每个run的文本内容和格式属性，
        包括字体、颜色、样式等信息。

        Args:
            paragraph: Word段落对象

        Returns:
            List[Dict]: runs信息列表，每个字典包含文本内容和格式属性
        """
        if paragraph is None:
            return []
            
        runs = []
        for run in paragraph.runs:
            if run.text:  # 只处理非空文本片段
                runs.append({
                    'text': run.text,
                    'run': run,  # 保存run对象引用
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name if run.font.name else None,
                    'font_size': run.font.size.pt if run.font.size else None,
                    'font_color': run.font.color.rgb if run.font.color else None,
                    'highlight_color': run.font.highlight_color if hasattr(run.font, 'highlight_color') else None
                })
        return runs
    
    def get_all_text(self) -> List[str]:
        """
        获取所有文本内容，用于处理

        遍历文档结构，提取所有段落和表格单元格中的文本内容，
        返回按顺序排列的文本列表。

        Returns:
            List[str]: 文本内容列表
        """
        texts = []
        for item in self.original_structure:
            if item['type'] == 'paragraph':
                texts.append(item['original_text'])
            elif item['type'] == 'table':
                for row in item['rows']:
                    for cell in row:
                        if cell['paragraphs']:
                            for para in cell['paragraphs']:
                                texts.append(para['original_text'])
                        else:
                            texts.append(cell['original_text'])
        return texts
    
    def update_text_content(self, processed_texts: List[str]) -> bool:
        """
        更新文档中的文本内容，严格保持原有格式不变
        
        Args:
            processed_texts: 处理后的文本列表
            
        Returns:
            bool: 更新成功返回True，否则返回False
        """
        try:
            text_index = 0
            
            # 更新段落文本
            for item in self.original_structure:
                if item['type'] == 'paragraph' and text_index < len(processed_texts):
                    # 保持段落的所有属性不变，只更新文本内容
                    self._update_paragraph_text(item['paragraph'], processed_texts[text_index])
                    text_index += 1
                    
                elif item['type'] == 'table':
                    # 更新表格中的文本
                    for row in item['rows']:
                        for cell in row:
                            if cell['paragraphs']:
                                for para in cell['paragraphs']:
                                    if text_index < len(processed_texts):
                                        self._update_paragraph_text(para['paragraph'], processed_texts[text_index])
                                        text_index += 1
                            else:
                                # 直接更新单元格文本
                                if text_index < len(processed_texts):
                                    self._update_cell_text(cell['cell'], processed_texts[text_index])
                                    text_index += 1
            
            return True
            
        except Exception as e:
            print(f"更新文本内容失败: {e}")
            return False
    
    def _update_paragraph_text(self, paragraph, new_text: str) -> None:
        """
        更新段落文本，保持所有格式属性不变

        该方法在段落级别更新文本内容，尽可能保留原有的格式属性。
        对于单个run的段落，直接替换文本并保持格式；
        对于多个run的段落，使用第一个run的格式应用于所有新文本。

        Args:
            paragraph: Word段落对象
            new_text: 新的文本内容
        """
        # 基于run级别更新文本，保持原有格式
        if not paragraph.runs:
            # 如果没有runs，直接添加
            paragraph.add_run(new_text)
            return
        
        # 计算原始文本与新文本的映射关系
        original_text = paragraph.text
        
        # 如果文本未改变，不做任何处理
        if original_text == new_text:
            return
        
        # 获取所有runs的格式信息
        runs_info = []
        for run in paragraph.runs:
            runs_info.append({
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'font_color': run.font.color.rgb if run.font.color.rgb else None,
            })
        
        # 清空runs，但保持段落属性
        for run in paragraph.runs:
            r = run._element
            r.getparent().remove(r)
        
        # 如果只有一个run，直接替换其文本并保持格式
        if len(runs_info) == 1:
            new_run = paragraph.add_run(new_text)
            run_info = runs_info[0]
            new_run.bold = run_info['bold']
            new_run.italic = run_info['italic']
            new_run.underline = run_info['underline']
            if run_info['font_name']:
                new_run.font.name = run_info['font_name']
            if run_info['font_size']:
                new_run.font.size = run_info['font_size']
            if run_info['font_color']:
                new_run.font.color.rgb = run_info['font_color']
        else:
            # 多个runs的情况：使用第一个run的格式应用于所有新文本
            run_info = runs_info[0]
            new_run = paragraph.add_run(new_text)
            new_run.bold = run_info['bold']
            new_run.italic = run_info['italic']
            new_run.underline = run_info['underline']
            if run_info['font_name']:
                new_run.font.name = run_info['font_name']
            if run_info['font_size']:
                new_run.font.size = run_info['font_size']
            if run_info['font_color']:
                new_run.font.color.rgb = run_info['font_color']
    
    def _update_cell_text(self, cell, new_text: str) -> None:
        """
        更新单元格文本，保持所有格式属性不变

        清空单元格中的旧文本，然后添加新文本内容。

        Args:
            cell: Word表格单元格对象
            new_text: 新的文本内容
        """
        # 清空单元格内容但保持格式
        cell.text = ''
        
        # 添加新的文本内容
        cell.paragraphs[0].add_run(new_text)
    
    def save_document(self, file_path: str) -> bool:
        """
        保存文档到指定路径
        
        执行完整的输入验证和错误处理。
        
        Args:
            file_path: 保存路径
            
        Returns:
            bool: 保存成功返回True
            
        Raises:
            DocumentSaveError: 当保存失败时抛出
            
        Examples:
            >>> processor = DocumentProcessor()
            >>> processor.load_document("input.docx")
            >>> processor.save_document("output.docx")
            True
        """
        # 输入验证
        if not file_path:
            self.logger.error("保存路径不能为空")
            raise ValueError("保存路径不能为空")
        
        if not self.document:
            self.logger.error("没有加载的文档，无法保存")
            raise DocumentSaveError("没有加载的文档，无法保存")
        
        file_path = Path(file_path)
        
        # 确保目标目录存在
        file_path.parent.mkdir(parents=True, exist_ok=True)
        
        try:
            self.logger.info(f"开始保存文档: {file_path}")
            self.document.save(str(file_path))
            self.logger.info(f"成功保存文档: {file_path}")
            return True
            
        except PermissionError as e:
            self.logger.error(f"权限不足，无法保存文件: {file_path}", exc_info=True)
            raise PermissionError(f"权限不足，无法保存文件: {file_path}") from e
        except OSError as e:
            self.logger.error(f"保存文档时发生OS错误: {file_path}", exc_info=True)
            raise DocumentSaveError(f"保存文档失败: {str(e)}") from e
        except Exception as e:
            self.logger.error(f"保存文档失败: {file_path}", exc_info=True)
            raise DocumentSaveError(f"保存文档失败: {str(e)}") from e
    
    def save_as_new_document(self, file_path: str, processed_texts: List[str]) -> bool:
        """
        创建新文档并保存处理后的文本，保持原有结构
        
        Args:
            file_path: 保存路径
            processed_texts: 处理后的文本列表
            
        Returns:
            bool: 保存成功返回True，否则返回False
        """
        try:
            # 首先更新原文档的内容
            if self.update_text_content(processed_texts):
                # 然后保存为新文件
                return self.save_document(file_path)
            return False
            
        except Exception as e:
            print(f"保存新文档失败: {e}")
            return False
    
    def get_document_info(self) -> Dict:
        """
        获取文档信息

        返回文档的基本统计信息，包括段落数、表格数等。

        Returns:
            Dict: 文档信息字典，包含paragraph_count、table_count等字段
        """
        if not self.document:
            return {}
        
        paragraph_count = len([item for item in self.original_structure if item['type'] == 'paragraph'])
        table_count = len([item for item in self.original_structure if item['type'] == 'table'])
        
        return {
            'paragraph_count': paragraph_count,
            'table_count': table_count,
            'total_text_items': len(self.original_structure)
        }